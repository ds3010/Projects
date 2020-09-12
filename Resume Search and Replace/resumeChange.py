import re
from docx import Document

doc = Document('Daniel_Seijas_Resume_short.docx')

print('#'*100)
print('Indicate how many changes do you want to make: ')
while(True):
    changes = input()
    try:
        changes = int(changes)
        break
    except ValueError:
        print('You must enter a number, please try again.')

for i in range(changes):
    
    counter = 0
    print('\n' + ('#'*100))
    print('Change #' + str(i+1) + '- Enter a string to search and replace: ')
    userInput = input()
    userInput = str(userInput)
    prev = re.compile(r'(%s)'%userInput)
    for p in doc.paragraphs:
        for r in p.runs:
            search = prev.search(r.text)
            if search != None:
                counter += 1
                print('\n' + ('#'*100))
                print('\n-->The string has been found here: ' + r.text)
                print('\nWould you like to replace it?(y/n): ')
                while True:
                    userChoice = input()
                    if userChoice != 'y' and userChoice != 'Y' and userChoice != 'n' and userChoice != 'N':
                        print('Wrong choice, please type \'y\' or \'n\':')
                    elif userChoice == 'n' or userChoice == 'N':
                        break
                    elif userChoice == 'y' or userChoice == 'Y':
                        ##### CODE TO REPLACE THE WORD #####
                        print('\nType a string of characters to replace it with: ')
                        userInput = input()
                        newString = prev.sub(userInput, r.text)
                        print('\nCOMPLETED! - String has been replaced to: ' + str(newString))
                        r.text = newString
                        ####################################
                        break
    #print('\nAmount of times found in paragraphs: ' + str(counter))
    
    counter = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for r in paragraph.runs:
                        search = prev.search(r.text)
                        if search != None:
                            counter += 1
                            print('\n' + ('#'*100))
                            print('\nThe string has been found here: ' + r.text)
                            print('\nWould you like to replace it?(y/n): ')
                            while True:
                                userChoice = input()
                                if userChoice != 'y' and userChoice != 'Y' and userChoice != 'n' and userChoice != 'N':
                                    print('\nWrong choice, please type \'y\' or \'n\':')
                                elif userChoice == 'n' or userChoice == 'N':
                                    break
                                elif userChoice == 'y' or userChoice == 'Y':
                                    ##### CODE TO REPLACE THE WORD #####
                                    print('\nType a string of characters to replace it with: ')
                                    userInput = input()
                                    newString = prev.sub(userInput, r.text)
                                    print('\nCOMPLETED! - String has been replaced to: ' + str(newString))
                                    r.text = newString
                                    ####################################
                                    break
    #print('\nAmount of times found in tables: ' + str(counter))     

print('\n' + ('#'*100))
print('\nTime to save your changes: (Type words like "SalesEnginner", "NetworkAdmin"):')
resType = input()
doc.save('DanielSeijas_' + str(resType) + '.docx')
print('\n' + ('#'*100))
print('CONGRATS!!!!!!' + 'A copy of your resume has been created with the name: DanielSeijas_' + str(resType) + '.docx')                   

#The following script prints just the text outside tables
# fullText = []
# for para in doc.paragraphs:
#     fullText.append(para.text)
# print('\n'.join(fullText))

#The following script prints just the text inside tables
# i = 0
# for table in doc.tables:
#     i += 1
#     print('\nTABLE ' + str(i) +'######:\n')
#     for row in table.rows:
#         for cell in row.cells:
#             for paragraph in cell.paragraphs:
#                 lines = paragraph.runs
#                 for i in range(len(lines)):
#                     print(type(lines[i].text))